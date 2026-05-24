"""Tests for resume PDF/DOCX parsing."""

from pathlib import Path

import fitz
import pytest
from docx import Document

from config.settings import Settings, get_settings
from src.ingestion.exceptions import (
    ExtractionFailedError,
    FileTooLargeError,
    UnsupportedFileTypeError,
)
from src.ingestion.resume_parser import _clean_text, parse_docx, parse_resume, validate_resume_file


@pytest.fixture
def tmp_resume_dir(tmp_path: Path) -> Path:
    """Temporary directory for generated resume files."""
    return tmp_path


def test_parse_pdf_extracts_text(tmp_resume_dir: Path) -> None:
    """PDF parsing returns embedded text."""
    pdf_path = tmp_resume_dir / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "Senior Python Engineer at Acme Corp. Built ML pipelines and APIs for five years.",
    )
    doc.save(pdf_path)
    doc.close()

    text = parse_resume(str(pdf_path))
    assert "Python Engineer" in text
    assert "Acme Corp" in text


def test_parse_docx_with_table(tmp_resume_dir: Path) -> None:
    """DOCX parsing includes table cell content."""
    docx_path = tmp_resume_dir / "sample.docx"
    document = Document()
    document.add_paragraph("Jane Developer")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "Advanced"
    document.save(docx_path)

    text = parse_docx(str(docx_path))
    assert "Jane Developer" in text
    assert "Python" in text
    assert "Advanced" in text


def test_clean_text_normalizes_whitespace_and_entities() -> None:
    """Cleaning collapses whitespace and decodes HTML entities."""
    raw = "  Hello&amp;World   \n\n\n  Foo   \n  "
    cleaned = _clean_text(raw + "x" * 60)
    assert "Hello&World" in cleaned
    assert "  " not in cleaned.replace("\n\n", "")


def test_unsupported_extension(tmp_resume_dir: Path) -> None:
    """Unsupported file types raise UnsupportedFileTypeError."""
    txt_path = tmp_resume_dir / "resume.txt"
    txt_path.write_text("plain text resume", encoding="utf-8")
    with pytest.raises(UnsupportedFileTypeError):
        validate_resume_file(str(txt_path))


def test_file_too_large(tmp_resume_dir: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """Oversized files raise FileTooLargeError."""
    get_settings.cache_clear()
    monkeypatch.setenv("RESUME_MAX_FILE_BYTES", "10")
    settings = Settings()

    docx_path = tmp_resume_dir / "big.docx"
    document = Document()
    document.add_paragraph("x" * 500)
    document.save(docx_path)

    with pytest.raises(FileTooLargeError):
        validate_resume_file(str(docx_path), settings=settings)


def test_empty_file_raises(tmp_resume_dir: Path) -> None:
    """Empty files fail validation."""
    empty = tmp_resume_dir / "empty.pdf"
    empty.write_bytes(b"")
    with pytest.raises(ExtractionFailedError):
        validate_resume_file(str(empty))


def test_missing_file_raises(tmp_resume_dir: Path) -> None:
    """Missing path raises FileNotFoundError."""
    with pytest.raises(FileNotFoundError):
        validate_resume_file(str(tmp_resume_dir / "missing.pdf"))


def test_too_short_text_raises() -> None:
    """Very short extracted text raises ExtractionFailedError."""
    with pytest.raises(ExtractionFailedError):
        _clean_text("short")
