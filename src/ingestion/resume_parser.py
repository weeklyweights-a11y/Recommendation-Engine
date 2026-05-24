"""Extract and clean text from PDF and DOCX resume files."""

from __future__ import annotations

import html
import logging
import re
import unicodedata
from pathlib import Path

import fitz
from docx import Document

from config.settings import Settings, get_settings
from src.ingestion.exceptions import (
    ExtractionFailedError,
    FileTooLargeError,
    UnsupportedFileTypeError,
)

logger = logging.getLogger(__name__)

_SUPPORTED_SUFFIXES = {".pdf", ".docx", ".doc"}


def validate_resume_file(file_path: str, settings: Settings | None = None) -> None:
    """Validate resume path, extension, and size before parsing."""
    path = Path(file_path)
    if not path.is_file():
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix not in _SUPPORTED_SUFFIXES:
        raise UnsupportedFileTypeError(f"Unsupported resume format: {suffix}")

    cfg = (settings or get_settings()).ingestion
    size = path.stat().st_size
    if size <= 0:
        raise ExtractionFailedError("Resume file is empty")
    if size > cfg.resume_max_file_bytes:
        raise FileTooLargeError(
            f"Resume file exceeds limit of {cfg.resume_max_file_bytes} bytes",
        )


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF using block-ordered reading and optional tables."""
    parts: list[str] = []
    with fitz.open(file_path) as doc:
        for page in doc:
            blocks = page.get_text("blocks")
            if blocks:
                sorted_blocks = sorted(blocks, key=lambda b: (round(b[1], 1), round(b[0], 1)))
                page_text = "\n".join(
                    str(block[4]).strip() for block in sorted_blocks if str(block[4]).strip()
                )
            else:
                page_text = page.get_text("text")

            table_parts: list[str] = []
            try:
                tables = page.find_tables()
                for table in tables:
                    rows = table.extract()
                    for row in rows:
                        cells = [str(cell or "").strip() for cell in row]
                        if any(cells):
                            table_parts.append(", ".join(cells))
            except Exception:  # noqa: BLE001 — table API varies by PyMuPDF version
                pass

            combined = "\n".join(part for part in [page_text, "\n".join(table_parts)] if part)
            if combined.strip():
                parts.append(combined)

    return "\n\n".join(parts)


def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file including table cells."""
    document = Document(file_path)
    segments: list[str] = [para.text.strip() for para in document.paragraphs if para.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                segments.append(" | ".join(cells))

    return "\n".join(segments)


def _clean_text(text: str, settings: Settings | None = None) -> str:
    """Normalize unicode, whitespace, and common PDF artifacts."""
    cfg = (settings or get_settings()).ingestion
    normalized = unicodedata.normalize("NFKD", text)
    normalized = html.unescape(normalized)
    normalized = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", normalized)

    lines = [line.strip() for line in normalized.splitlines()]
    non_empty = [line for line in lines if line]
    if len(non_empty) >= 3:
        counts: dict[str, int] = {}
        for line in non_empty:
            counts[line] = counts.get(line, 0) + 1
        repeated = {line for line, count in counts.items() if count >= 3 and len(line) < 80}
        if repeated:
            non_empty = [line for line in non_empty if line not in repeated]

    cleaned = "\n\n".join(non_empty)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = cleaned.strip()

    if len(cleaned) < cfg.resume_min_text_chars:
        raise ExtractionFailedError("Could not extract meaningful text from resume")

    if len(cleaned) > cfg.resume_max_text_chars:
        logger.warning(
            "Resume text truncated from %s to %s characters",
            len(cleaned),
            cfg.resume_max_text_chars,
        )
        cleaned = cleaned[: cfg.resume_max_text_chars]

    return cleaned


def parse_resume(file_path: str, settings: Settings | None = None) -> str:
    """Detect format, parse, and return cleaned resume text."""
    validate_resume_file(file_path, settings=settings)
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        raw = parse_pdf(file_path)
    elif suffix in {".docx", ".doc"}:
        raw = parse_docx(file_path)
    else:
        raise UnsupportedFileTypeError(f"Unsupported resume format: {suffix}")

    return _clean_text(raw, settings=settings)
